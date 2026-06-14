import base64
import io
import logging
from typing import Any

import fitz  # PyMuPDF — robust text extraction + page rendering for vision OCR.
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx import Document as DocxDocument

from app.config import Settings
from app.services.vector_store import add_chunks, delete_document_chunks, get_vector_store

logger = logging.getLogger(__name__)

PDF_INGEST_VERSION = "pymupdf-gemini-v1"

# Pages with fewer characters than this are treated as image-only and sent to Gemini vision.
_MIN_PAGE_TEXT_CHARS = 12
# Cap vision OCR so large scanned files cannot exhaust the free-tier instance.
_MAX_VISION_PAGES = 20
_VISION_DPI = 150


def _split_documents(
    documents: list[Document],
    settings: Settings,
) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_documents(documents)


def _gemini_ocr_page(settings: Settings, page: fitz.Page) -> str:
    """Extract text from an image-only PDF page via Gemini vision."""
    if not settings.google_api_key.strip():
        return ""

    pix = page.get_pixmap(dpi=_VISION_DPI)
    b64 = base64.standard_b64encode(pix.tobytes("png")).decode("ascii")

    llm = ChatGoogleGenerativeAI(
        google_api_key=settings.google_api_key,
        model=settings.gemini_model,
        temperature=0,
    )
    message = HumanMessage(
        content=[
            {
                "type": "text",
                "text": (
                    "Extract every word of visible text from this document page. "
                    "Return plain text only, preserving useful line breaks. "
                    "If there is no readable text, return an empty string."
                ),
            },
            {"type": "image_url", "image_url": f"data:image/png;base64,{b64}"},
        ]
    )
    result = llm.invoke([message])
    return (result.content if hasattr(result, "content") else str(result)).strip()


def _pdf_to_documents(data: bytes, settings: Settings) -> list[Document]:
    """Extract text with PyMuPDF; image-only pages fall back to Gemini vision OCR."""
    docs: list[Document] = []
    vision_pages: list[int] = []

    with fitz.open(stream=data, filetype="pdf") as pdf:
        for page_num in range(1, pdf.page_count + 1):
            page = pdf[page_num - 1]
            text = (page.get_text("text") or "").strip()
            if len(text) >= _MIN_PAGE_TEXT_CHARS:
                docs.append(Document(page_content=text, metadata={"page": page_num}))
            else:
                vision_pages.append(page_num)

        if vision_pages and settings.google_api_key.strip():
            for page_num in vision_pages[:_MAX_VISION_PAGES]:
                try:
                    ocr_text = _gemini_ocr_page(settings, pdf[page_num - 1])
                except Exception:
                    logger.exception("Gemini vision OCR failed for page %s", page_num)
                    continue
                if len(ocr_text) >= _MIN_PAGE_TEXT_CHARS:
                    docs.append(Document(page_content=ocr_text, metadata={"page": page_num}))
        elif vision_pages:
            logger.warning(
                "%d page(s) had no extractable text and Gemini is not configured.",
                len(vision_pages),
            )

    docs.sort(key=lambda d: int(d.metadata.get("page", 0)))
    return docs


def _docx_to_documents(data: bytes) -> list[Document]:
    doc = DocxDocument(io.BytesIO(data))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        return []
    return [Document(page_content="\n\n".join(paragraphs), metadata={"page": 1})]


def ingest_docx_bytes(
    settings: Settings,
    document_id: str,
    data: bytes,
    title: str | None = None,
) -> dict[str, Any]:
    """Index Word (.docx) documents for RAG."""
    raw_docs = _docx_to_documents(data)
    if not raw_docs:
        raise ValueError("No extractable text found in Word document")

    chunks = _split_documents(raw_docs, settings)
    store = get_vector_store(settings)
    delete_document_chunks(store, document_id)
    count = add_chunks(store, document_id, title, chunks)

    return {
        "document_id": document_id,
        "status": "indexed",
        "chunk_count": count,
        "title": title,
    }


def ingest_pdf_bytes(
    settings: Settings,
    document_id: str,
    data: bytes,
    title: str | None = None,
) -> dict[str, Any]:
    """Blueprint Step 4: chunk PDF → embed → ChromaDB."""
    raw_docs = _pdf_to_documents(data, settings)
    if not raw_docs:
        raise ValueError(
            "We couldn't read any text from this PDF. It may be an image-only or "
            "password-protected file. Try a text-based PDF or a .docx version."
        )

    chunks = _split_documents(raw_docs, settings)
    store = get_vector_store(settings)
    delete_document_chunks(store, document_id)
    count = add_chunks(store, document_id, title, chunks)

    return {
        "document_id": document_id,
        "status": "indexed",
        "chunk_count": count,
        "title": title,
    }


def ingest_text(
    settings: Settings,
    document_id: str,
    text: str,
    title: str | None = None,
) -> dict[str, Any]:
    """Index plain text (Word export, pasted content, tests)."""
    raw = [Document(page_content=text.strip(), metadata={"page": 1})]
    chunks = _split_documents(raw, settings)
    store = get_vector_store(settings)
    delete_document_chunks(store, document_id)
    count = add_chunks(store, document_id, title, chunks)

    return {
        "document_id": document_id,
        "status": "indexed",
        "chunk_count": count,
        "title": title,
    }
